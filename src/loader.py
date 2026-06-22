import json
import gzip
import logging
import os
import docx
from typing import List, Dict, Any, Tuple

logger = logging.getLogger(__name__)

def load_candidates(candidates_path: str) -> List[Dict[str, Any]]:
    """
    Load candidate data from a JSONL or JSONL.GZ file.
    """
    if not os.path.exists(candidates_path):
        raise FileNotFoundError(f"Candidates file not found at: {candidates_path}")
        
    candidates = []
    logger.info(f"Loading candidates from {candidates_path}...")
    
    # Handle gzip
    if candidates_path.endswith(".gz"):
        open_func = gzip.open
        mode = "rt"
    else:
        open_func = open
        mode = "r"
        
    with open_func(candidates_path, mode, encoding="utf-8") as f:
        for line_num, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                candidate_data = json.loads(line)
                if "id" not in candidate_data and "candidate_id" in candidate_data:
                    candidate_data["id"] = candidate_data["candidate_id"]
                elif "id" not in candidate_data:
                    candidate_data["id"] = f"cand_{line_num}"
                candidates.append(candidate_data)
            except json.JSONDecodeError as e:
                logger.warning(f"Skipping line {line_num} due to JSON decode error: {e}")
                
    logger.info(f"Successfully loaded {len(candidates)} candidates.")
    return candidates


def load_jd_text(jd_path: str) -> str:
    """
    Read text and tables from docx file and return it as a single string.
    """
    if not os.path.exists(jd_path):
        raise FileNotFoundError(f"JD file not found at: {jd_path}")
        
    doc = docx.Document(jd_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in table_texts:
                    table_texts.append(text)
    return "\n".join(paragraphs + table_texts)


def load_all(candidates_path: str, jd_path: str) -> Tuple[List[Dict[str, Any]], str]:
    """
    Load all candidate data (handles jsonl and jsonl.gz) and the job description text.
    """
    candidates = load_candidates(candidates_path)
    jd_text = load_jd_text(jd_path)
    return candidates, jd_text
