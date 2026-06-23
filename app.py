import os
# Disable file watcher to prevent PyTorch classes reflection crashes
os.environ["STREAMLIT_SERVER_ENABLE_FILE_WATCHER"] = "false"
try:
    import torch
    torch.classes.__path__ = []
except ImportError:
    pass

import streamlit as st
import time
import json
import docx
import pandas as pd
import numpy as np
import plotly.express as px
import textwrap

# Import pipeline modules
import src.jd_parser as jd_parser
import src.feature_extractor as feature_extractor
import src.scorer as scorer
import src.honeypot_detector as honeypot_detector
import src.reasoning_generator as reasoning_generator

# Streamlit App Configuration
st.set_page_config(
    page_title="Redrob Ranker",
    page_icon="🏆",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom styling injection for a premium SaaS look
st.markdown(
    """
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');
        
        /* Font overrides */
        html, body, [class*="css"], .stMarkdown, .stText, p, span, h1, h2, h3, h4, h5, h6, label {
            font-family: 'Outfit', sans-serif !important;
        }
        
        /* Force global background to Dark Navy */
        .stApp, div[data-testid="stAppViewContainer"], .main, [data-testid="stAppViewContainer"] {
            background-color: #0F1B2D !important;
            color: #FFFFFF !important;
        }
        
        /* Transparent Header */
        div[data-testid="stHeader"] {
            background-color: transparent !important;
        }
        
        /* Sidebar styling */
        section[data-testid="stSidebar"] {
            background-color: #0A1320 !important;
            border-right: 1px solid #263954 !important;
        }
        section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p {
            color: #FFFFFF !important;
        }
        
        /* Premium Gradient Header */
        .gradient-text {
            background: linear-gradient(135deg, #00C9A7 0%, #00F2FE 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-size: 42px;
            font-weight: 700;
            margin-bottom: 0.1rem;
            margin-top: -30px;
            letter-spacing: -0.02em;
        }
        
        /* Subtitle */
        .subtitle-text {
            color: #94A3B8;
            font-size: 18px;
            font-weight: 400;
            margin-bottom: 2rem;
        }
        
        /* Card design */
        .metric-card {
            background-color: #16253B;
            border: 1px solid #263954;
            border-radius: 16px;
            padding: 20px;
            box-shadow: 0 4px 20px rgba(0, 0, 0, 0.15);
            transition: all 0.3s ease;
            text-align: center;
        }
        .metric-card:hover {
            transform: translateY(-4px);
            border-color: #00C9A7;
            box-shadow: 0 8px 30px rgba(0, 201, 167, 0.15);
        }
        .metric-val {
            font-size: 32px;
            font-weight: 700;
            color: #00C9A7;
            margin-bottom: 2px;
        }
        .metric-label {
            font-size: 12px;
            color: #94A3B8;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }
        
        /* Custom Progress Bars for Signal Analysis */
        .signal-row {
            margin-bottom: 16px;
        }
        .signal-header {
            display: flex;
            justify-content: space-between;
            margin-bottom: 6px;
            font-size: 14px;
            font-weight: 600;
            color: #E2E8F0;
        }
        .signal-name {
            font-family: monospace;
            background-color: #0F1B2D;
            padding: 2px 8px;
            border-radius: 4px;
            border: 1px solid #263954;
            font-size: 12px;
        }
        .signal-bar-bg {
            background-color: #0A1320;
            height: 10px;
            border-radius: 5px;
            overflow: hidden;
            border: 1px solid #1E2D44;
        }
        .signal-bar-fill {
            background: linear-gradient(90deg, #00C9A7 0%, #00F2FE 100%);
            height: 100%;
            border-radius: 5px;
        }
        
        /* Premium custom buttons */
        .stButton button {
            background: linear-gradient(135deg, #00C9A7 0%, #009E86 100%) !important;
            color: #0F1B2D !important;
            font-weight: 700 !important;
            border: none !important;
            border-radius: 12px !important;
            padding: 10px 20px !important;
            transition: all 0.3s ease !important;
            box-shadow: 0 4px 15px rgba(0, 201, 167, 0.25) !important;
            width: 100% !important;
        }
        .stButton button:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 6px 20px rgba(0, 201, 167, 0.35) !important;
            color: #0F1B2D !important;
        }
        
        /* Style download buttons specifically to have a dark outline and dark text/label */
        div[data-testid="stDownloadButton"] button, .stDownloadButton button {
            background-color: #FFFFFF !important;
            color: #0F1B2D !important;
            border: 2px solid #0F1B2D !important;
            font-weight: 700 !important;
            border-radius: 12px !important;
            padding: 10px 20px !important;
            transition: all 0.3s ease !important;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1) !important;
            width: 100% !important;
        }
        div[data-testid="stDownloadButton"] button:hover, .stDownloadButton button:hover {
            background-color: #0F1B2D !important;
            color: #FFFFFF !important;
            border: 2px solid #00C9A7 !important;
            transform: translateY(-2px) !important;
            box-shadow: 0 6px 20px rgba(0, 201, 167, 0.3) !important;
        }
        
        /* Tabs overrides */
        .stTabs [data-baseweb="tab-list"] {
            gap: 12px !important;
            background-color: #16253B !important;
            padding: 8px !important;
            border-radius: 16px !important;
            border: 1px solid #263954 !important;
            margin-bottom: 20px !important;
        }
        .stTabs [data-baseweb="tab"] {
            color: #94A3B8 !important;
            background-color: transparent !important;
            border-radius: 10px !important;
            padding: 10px 20px !important;
            font-weight: 600 !important;
            border: none !important;
            transition: all 0.2s ease !important;
        }
        .stTabs [data-baseweb="tab"]:hover {
            color: #00C9A7 !important;
            background-color: rgba(0, 201, 167, 0.05) !important;
        }
        .stTabs [aria-selected="true"] {
            background-color: #00C9A7 !important;
            color: #0F1B2D !important;
            box-shadow: 0 4px 12px rgba(0, 201, 167, 0.2) !important;
        }
        
        /* Table headers styling */
        th {
            background-color: #16253B !important;
            color: #00C9A7 !important;
            font-weight: 700 !important;
            border-bottom: 2px solid #263954 !important;
        }
        td {
            border-bottom: 1px solid #1E2D44 !important;
        }
        
        /* File uploader visual style overrides */
        [data-testid="stFileUploader"] {
            border: 1px dashed #263954 !important;
            border-radius: 12px !important;
            background-color: #16253B !important;
            padding: 15px !important;
        }
        [data-testid="stFileUploader"] section {
            padding: 0 !important;
            background-color: #16253B !important;
            color: #FFFFFF !important;
        }
        [data-testid="stFileUploaderDropzone"] {
            background-color: #16253B !important;
            border: 1px solid #263954 !important;
            border-radius: 8px !important;
            color: #FFFFFF !important;
        }
        [data-testid="stFileUploaderDropzone"] > div {
            color: #E2E8F0 !important;
        }
        /* Style text and small limit labels in dropzone to be crisp and clear */
        [data-testid="stFileUploaderDropzone"] p,
        [data-testid="stFileUploaderDropzone"] span {
            color: #FFFFFF !important;
            font-size: 15px !important;
            font-weight: 600 !important;
        }
        [data-testid="stFileUploaderDropzone"] small {
            color: #94A3B8 !important;
            font-size: 13px !important;
            font-weight: 500 !important;
            display: block !important;
            margin-top: 4px !important;
            opacity: 1 !important;
        }
        [data-testid="stFileUploaderDropzone"] button {
            background-color: #263954 !important;
            color: #FFFFFF !important;
            border: 1px solid #3B5374 !important;
            border-radius: 8px !important;
            padding: 8px 16px !important;
            font-size: 14px !important;
            font-weight: 600 !important;
            transition: all 0.2s ease !important;
        }
        [data-testid="stFileUploaderDropzone"] button:hover {
            background-color: #00C9A7 !important;
            color: #0F1B2D !important;
            border-color: #00C9A7 !important;
        }
        [data-testid="stFileUploaderFileName"], [data-testid="stFileUploaderDeleteBtn"] {
            color: #FFFFFF !important;
        }
        
        /* Info/Success metrics box styling */
        .stAlert {
            background-color: #16253B !important;
            border: 1px solid #263954 !important;
            color: #E2E8F0 !important;
            border-radius: 12px !important;
        }
    </style>
    """,
    unsafe_allow_html=True
)

# Cache the SentenceTransformer model to load it once
@st.cache_resource
def load_sentence_transformer_model():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer("all-MiniLM-L6-v2")

# Parser for candidates (.json or .jsonl)
def parse_uploaded_candidates(uploaded_file):
    content = uploaded_file.getvalue().decode("utf-8").strip()
    if content.startswith("["):
        try:
            return json.loads(content)
        except Exception as e:
            st.error(f"Error parsing JSON list: {e}")
            return None
    else:
        candidates = []
        for line_num, line in enumerate(content.splitlines(), 1):
            line = line.strip()
            if not line:
                continue
            try:
                candidate_data = json.loads(line)
                if "id" not in candidate_data and "candidate_id" in candidate_data:
                    candidate_data["id"] = candidate_data["candidate_id"]
                elif "id" not in candidate_data:
                    candidate_data["id"] = f"cand_{line_num}"
                candidates.append(candidate_data)
            except Exception as e:
                st.warning(f"Skipping line {line_num} due to parse error: {e}")
        return candidates

# Parser for JD docx
def parse_uploaded_jd(uploaded_file):
    try:
        doc = docx.Document(uploaded_file)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in table_texts:
                        table_texts.append(text)
        return "\n".join(paragraphs + table_texts)
    except Exception as e:
        st.error(f"Error parsing JD docx file: {e}")
        return None

# Core pipeline execution logic for the sandbox app
def run_pipeline(candidates, jd_text):
    # Pre-load/inject the cached model into scorer singleton
    model = load_sentence_transformer_model()
    scorer._model = model
    
    # 1. Parse JD
    jd_profile = jd_parser.parse_jd(jd_text)
    
    # === STAGE 1 — Fast pre-filter (Rule-based) ===
    features_list = []
    for cand in candidates:
        features_list.append(feature_extractor.extract_features(cand))
        
    stage1_survivors = []
    
    rejection_titles = [
        "Marketing", "HR Manager", "Sales", "Graphic Designer", 
        "Content Writer", "Finance", "Recruiter", "Teacher", "Doctor", "Lawyer"
    ]
    rejection_titles_lower = [t.lower() for t in rejection_titles]
    
    # Quick lookup dictionary for candidates
    cand_lookup = {str(c.get("id", c.get("candidate_id", ""))): c for c in candidates}
    
    for feat in features_list:
        cand_id = feat["candidate_id"]
        cand = cand_lookup.get(cand_id)
        if not cand:
            continue
            
        # Filter 1: years_experience
        years_exp = feat.get("years_experience", 0.0)
        if years_exp < 2.0 or years_exp > 20.0:
            continue
            
        # Filter 2: consulting_only
        if feat.get("consulting_only", False):
            continue
            
        # Filter 3: inactive > 120 days
        if feat.get("days_since_active", 9999) > 120:
            continue
            
        # Filter 4: keyword_stuffer
        if feat.get("keyword_stuffer", False):
            continue
            
        # Filter 5: impossible_experience (honeypot)
        if feat.get("impossible_experience", False):
            continue
            
        # Filter 6: country and relocate
        country_lower = feat.get("country", "").strip().lower()
        willing = feat.get("willing_to_relocate", False)
        if country_lower != "india" and not willing:
            continue
            
        # Filter 7: title mismatch
        title_lower = feat.get("current_title", "").strip().lower()
        if any(rt in title_lower for rt in rejection_titles_lower):
            continue
            
        # Compute hard match score
        h_score = scorer.fast_hard_score(feat)
        
        stage1_survivors.append({
            "candidate": cand,
            "features": feat,
            "hard_match_score": h_score
        })
        
    if len(stage1_survivors) == 0:
        raise ValueError("Validation Error: 0 candidates remain after Stage 1 pre-filters.")
        
    # Sort and keep top 3000 survivors for Stage 2
    stage1_survivors.sort(key=lambda x: str(x["features"]["candidate_id"]))
    stage1_survivors.sort(key=lambda x: float(x["hard_match_score"]), reverse=True)
    top_3000 = stage1_survivors[:3000]
    
    # === STAGE 2 — Deep scoring ===
    top_3000_features = [item["features"] for item in top_3000]
    scores_list = scorer.batch_score_all(top_3000_features, jd_profile)
    
    scored_candidates = []
    for item, score in zip(top_3000, scores_list):
        cand = item["candidate"]
        feat = item["features"]
        
        is_hard_hp, confidence = honeypot_detector.is_honeypot(cand)
        final_score = round(score * 0.05, 4) if is_hard_hp else score
        
        scored_candidates.append({
            "candidate_id": feat["candidate_id"],
            "score": final_score,
            "features": feat,
            "candidate": cand,
            "is_hard_hp": is_hard_hp
        })
        
    scored_candidates.sort(key=lambda x: str(x["candidate_id"]))
    scored_candidates.sort(key=lambda x: float(x["score"]), reverse=True)
    
    # Slice to top 100 for output
    ranked_results = scored_candidates[:100]
    
    # === STAGE 3 — Reasoning ===
    seen_reasonings = set()
    for rank, item in enumerate(ranked_results, 1):
        reason = reasoning_generator.generate_reasoning(item["features"], item["score"], rank)
        if reason in seen_reasonings:
            suffix = f" [{item['candidate_id']}]"
            if len(reason) + len(suffix) <= 200:
                reason = reason.rstrip(".") + suffix + "."
            else:
                reason = reason[:200 - len(suffix)] + suffix
        seen_reasonings.add(reason)
        
        if item["is_hard_hp"]:
            reason = f"FILTERED (hard): Flagged honeypot. " + reason
            if len(reason) > 200:
                reason = reason[:197] + "..."
        item["reasoning"] = reason
        
    return {
        "ranked_results": ranked_results,
        "stage1_survivors": stage1_survivors,
        "all_candidates": candidates,
        "total_processed": len(features_list),
        "passed_prefilter": len(stage1_survivors)
    }

# App UI Header
st.markdown('<div class="gradient-text">Redrob Ranker — AI Candidate Discovery</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle-text">Ranks candidates the way a great recruiter would</div>', unsafe_allow_html=True)

# Sidebar Upload Controls
st.sidebar.header("Upload your files")
uploaded_jd = st.sidebar.file_uploader("Job Description (.docx)", type=["docx"])
uploaded_candidates = st.sidebar.file_uploader("Candidates (.json or .jsonl)", type=["json", "jsonl"])
run_button = st.sidebar.button("🚀 Run Ranking")

st.sidebar.markdown("---")
st.sidebar.subheader("About")
st.sidebar.caption(
    "2-stage pipeline: rule-based pre-filter → "
    "semantic embedding scoring. No API calls. CPU only."
)

# Reset state if new files uploaded
if "last_jd" not in st.session_state or st.session_state.last_jd != uploaded_jd or \
   "last_candidates" not in st.session_state or st.session_state.last_candidates != uploaded_candidates:
    if "results" in st.session_state:
        del st.session_state.results
    if "stats" in st.session_state:
        del st.session_state.stats
    st.session_state.last_jd = uploaded_jd
    st.session_state.last_candidates = uploaded_candidates

# Pipeline run trigger
if run_button:
    if not uploaded_jd:
        st.error("Please upload a Job Description (.docx) file first.")
    elif not uploaded_candidates:
        st.error("Please upload Candidates (.json or .jsonl) file first.")
    else:
        with st.spinner("Running pipeline..."):
            t0 = time.time()
            # 1. Parse uploads
            jd_text = parse_uploaded_jd(uploaded_jd)
            candidates = parse_uploaded_candidates(uploaded_candidates)
            
            if jd_text and candidates:
                try:
                    # 2. Run ranking
                    results_dict = run_pipeline(candidates, jd_text)
                    t_elapsed = time.time() - t0
                    
                    # Store in session state
                    st.session_state.results = results_dict["ranked_results"]
                    st.session_state.all_candidates = results_dict["all_candidates"]
                    st.session_state.stage1_survivors = results_dict["stage1_survivors"]
                    st.session_state.stats = {
                        "total_processed": results_dict["total_processed"],
                        "passed_prefilter": results_dict["passed_prefilter"],
                        "runtime": t_elapsed
                    }
                    st.success("Pipeline executed successfully!")
                except Exception as e:
                    st.error(f"Pipeline execution failed: {e}")

# Render Results UI if cached in session state
if "results" in st.session_state and "stats" in st.session_state:
    ranked_results = st.session_state.results
    stats = st.session_state.stats
    
    # Build a complete results DataFrame
    rows = []
    for idx, item in enumerate(ranked_results, 1):
        rows.append({
            "Rank": idx,
            "Candidate ID": item["candidate_id"],
            "Score": item["score"],
            "Current Title": item["features"].get("current_title", "N/A"),
            "Years Exp": item["features"].get("years_experience", 0.0),
            "Location": item["features"].get("location", "N/A"),
            "Reasoning": item["reasoning"]
        })
    df_full = pd.DataFrame(rows)
    df_top10 = df_full.head(10)
    
    # Tabs layout
    tab1, tab2, tab3, tab4 = st.tabs([
        "🏆 Top 10 Results", 
        "📊 Score Distribution", 
        "🔍 Pipeline Stats", 
        "📥 Download"
    ])
    
    # Tab 1 — Styled Top 10 results
    with tab1:
        st.subheader("Top 10 Ranked Candidates")
        # Greens color scale gradient for the Score column
        styled_df = df_top10.style.background_gradient(subset=["Score"], cmap="Greens")
        st.dataframe(styled_df, use_container_width=True, hide_index=True)
        
    # Tab 2 — Score Distribution Chart
    with tab2:
        st.subheader("Candidate Score Distribution")
        # Beautiful custom Plotly chart matching theme
        fig = px.bar(
            df_full, 
            x="Rank", 
            y="Score",
            color="Score",
            color_continuous_scale=[[0, '#16253B'], [1, '#00C9A7']],
            labels={"Rank": "Rank", "Score": "Score"},
            title="Candidate Scores Sorted Descending"
        )
        fig.update_layout(
            plot_bgcolor="rgba(0,0,0,0)",
            paper_bgcolor="rgba(0,0,0,0)",
            font_family="Outfit",
            font_color="#E2E8F0",
            title_font_size=20,
            title_font_color="#FFFFFF",
            xaxis=dict(gridcolor="#1E2D44", showline=True, linecolor="#263954"),
            yaxis=dict(gridcolor="#1E2D44", showline=True, linecolor="#263954"),
            coloraxis_showscale=False,
            margin=dict(l=40, r=40, t=60, b=40)
        )
        st.plotly_chart(fig, use_container_width=True)
        
    # Tab 3 — Pipeline Metrics & Signals Table
    with tab3:
        # Custom visual cards for metrics
        st.markdown(
            textwrap.dedent(
                f"""
                <div style="display: flex; gap: 20px; margin-bottom: 30px;">
                    <div class="metric-card" style="flex: 1;">
                        <div class="metric-val">{stats['total_processed']}</div>
                        <div class="metric-label">Total Candidates Processed</div>
                    </div>
                    <div class="metric-card" style="flex: 1;">
                        <div class="metric-val">{stats['passed_prefilter']}</div>
                        <div class="metric-label">Passed Pre-filter (Stage 1)</div>
                    </div>
                    <div class="metric-card" style="flex: 1;">
                        <div class="metric-val">{stats['runtime']:.3f} s</div>
                        <div class="metric-label">Runtime (Seconds)</div>
                    </div>
                </div>
                """
            ),
            unsafe_allow_html=True
        )
        
        # Clickable drill-down actions
        st.markdown("##### 🔍 Drill-Down Interactive Inspector")
        if "selected_view" not in st.session_state:
            st.session_state.selected_view = "none"
            
        col_act1, col_act2, col_act3 = st.columns([1.2, 1.2, 2.1])
        with col_act1:
            if st.button("📋 Inspect All Uploaded", key="btn_view_all"):
                st.session_state.selected_view = "all_uploaded"
        with col_act2:
            if st.button("🎯 Inspect Pre-filtered", key="btn_view_passed"):
                st.session_state.selected_view = "passed_prefilter"
        with col_act3:
            if st.session_state.selected_view != "none":
                if st.button("❌ Hide Candidate Lists", key="btn_hide_all"):
                    st.session_state.selected_view = "none"
                    
        # Render the selected interactive lists
        if st.session_state.selected_view == "all_uploaded" and "all_candidates" in st.session_state:
            st.markdown("#### 📋 List of All Uploaded Candidates")
            all_cands = st.session_state.all_candidates
            all_df = pd.DataFrame([{
                "ID": c.get("id", c.get("candidate_id", "")),
                "Title": c.get("profile", {}).get("current_title", "N/A"),
                "Years Exp": c.get("profile", {}).get("years_of_experience", 0.0),
                "Location": c.get("profile", {}).get("location", "N/A"),
                "Country": c.get("profile", {}).get("country", "N/A")
            } for c in all_cands])
            st.dataframe(all_df, use_container_width=True, hide_index=True)
            
        elif st.session_state.selected_view == "passed_prefilter" and "stage1_survivors" in st.session_state:
            st.markdown("#### 🎯 List of Pre-filter Survivors (Stage 1)")
            survivors = st.session_state.stage1_survivors
            surv_df = pd.DataFrame([{
                "ID": s["features"]["candidate_id"],
                "Title": s["features"].get("current_title", "N/A"),
                "Years Exp": s["features"].get("years_experience", 0.0),
                "Location": s["features"].get("location", "N/A"),
                "Hard Match Score": s["hard_match_score"]
            } for s in survivors])
            st.dataframe(surv_df, use_container_width=True, hide_index=True)
            
        st.markdown("---")
        st.subheader("Signals Breakdown (Top 10)")
        top_10_features = [item["features"] for item in ranked_results[:10]]
        n_top_10 = len(top_10_features)
        
        if n_top_10 > 0:
            pct_otw = (sum(1 for f in top_10_features if f.get("open_to_work", False)) / n_top_10) * 100.0
            pct_active = (sum(1 for f in top_10_features if f.get("days_since_active", 9999) <= 90) / n_top_10) * 100.0
            pct_emb = (sum(1 for f in top_10_features if f.get("has_embeddings_experience", False)) / n_top_10) * 100.0
            pct_india = (sum(1 for f in top_10_features if f.get("location_fit") in ["exact", "tier1_india", "india_other"]) / n_top_10) * 100.0
            pct_notice = (sum(1 for f in top_10_features if f.get("notice_days", 30) <= 30) / n_top_10) * 100.0
        else:
            pct_otw = pct_active = pct_emb = pct_india = pct_notice = 0.0
            
        # Draw custom visual card for breakdown of signals with HTML progress bars
        st.markdown(
            textwrap.dedent(
                f"""
                <div class="metric-card" style="text-align: left; max-width: 700px; margin-top: 15px;">
                    <h4 style="margin-top: 0; color: #FFFFFF; font-weight: 600; font-size: 16px; margin-bottom: 20px;">Top Candidate Characteristics</h4>
                    
                    <div class="signal-row">
                        <div class="signal-header">
                            <span class="signal-name">open_to_work</span>
                            <span style="color: #00C9A7;">{pct_otw:.1f}%</span>
                        </div>
                        <div class="signal-bar-bg">
                            <div class="signal-bar-fill" style="width: {pct_otw}%;"></div>
                        </div>
                    </div>
                    
                    <div class="signal-row">
                        <div class="signal-header">
                            <span class="signal-name">active last 90d</span>
                            <span style="color: #00C9A7;">{pct_active:.1f}%</span>
                        </div>
                        <div class="signal-bar-bg">
                            <div class="signal-bar-fill" style="width: {pct_active}%;"></div>
                        </div>
                    </div>
                    
                    <div class="signal-row">
                        <div class="signal-header">
                            <span class="signal-name">has_embeddings_exp</span>
                            <span style="color: #00C9A7;">{pct_emb:.1f}%</span>
                        </div>
                        <div class="signal-bar-bg">
                            <div class="signal-bar-fill" style="width: {pct_emb}%;"></div>
                        </div>
                    </div>
                    
                    <div class="signal-row">
                        <div class="signal-header">
                            <span class="signal-name">india_based</span>
                            <span style="color: #00C9A7;">{pct_india:.1f}%</span>
                        </div>
                        <div class="signal-bar-bg">
                            <div class="signal-bar-fill" style="width: {pct_india}%;"></div>
                        </div>
                    </div>
                    
                    <div class="signal-row">
                        <div class="signal-header">
                            <span class="signal-name">notice &le; 30d</span>
                            <span style="color: #00C9A7;">{pct_notice:.1f}%</span>
                        </div>
                        <div class="signal-bar-bg">
                            <div class="signal-bar-fill" style="width: {pct_notice}%;"></div>
                        </div>
                    </div>
                </div>
                """
            ),
            unsafe_allow_html=True
        )
        
    # Tab 4 — Download Submission CSV
    with tab4:
        st.subheader("View and Download Complete Ranked Dataset")
        st.dataframe(df_full, use_container_width=True, hide_index=True)
        
        # Prepare CSV format matching submission specification
        csv_export = df_full[["Candidate ID", "Rank", "Score", "Reasoning"]].rename(columns={
            "Candidate ID": "candidate_id",
            "Rank": "rank",
            "Score": "score",
            "Reasoning": "reasoning"
        }).to_csv(index=False)
        
        st.download_button(
            label="Download submission.csv",
            data=csv_export,
            file_name="submission.csv",
            mime="text/csv"
        )
else:
    st.info("Upload files in the sidebar and click 'Run Ranking' to see candidate discovery analytics.")

if __name__ == "__main__":
    pass
