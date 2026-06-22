import argparse
import sys
import json
import logging
import time
import docx
import numpy as np

import src.loader as loader
import src.jd_parser as jd_parser
import src.honeypot_detector as honeypot_detector
import src.feature_extractor as feature_extractor
import src.scorer as scorer
import src.reasoning_generator as reasoning_generator
import src.output_writer as output_writer

def load_sample_candidates(path: str) -> list:
    """
    Load candidates from JSON (either JSON list or JSONL format).
    """
    with open(path, "r", encoding="utf-8") as f:
        content = f.read().strip()
    if content.startswith("["):
        return json.loads(content)
    else:
        candidates = []
        for line in content.splitlines():
            if line.strip():
                candidates.append(json.loads(line))
        return candidates

def main():
    logging.basicConfig(level=logging.ERROR)  # Suppress debug logs during test run
    
    parser = argparse.ArgumentParser(description="Local Validation and Testing Script")
    parser.add_argument("--sample", required=True, help="Path to sample candidates JSON file")
    parser.add_argument("--jd", required=True, help="Path to job description DOCX file")
    args = parser.parse_args()

    print(f"Loading sample candidates from: {args.sample}")
    print(f"Loading job description from: {args.jd}")

    try:
        # 1. Load data
        candidates = load_sample_candidates(args.sample)
        print(f"Loaded {len(candidates)} candidates.")
        
        # Load JD
        doc = docx.Document(args.jd)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in table_texts:
                        table_texts.append(text)
        jd_text = "\n".join(paragraphs + table_texts)
        
        # 2. Parse JD
        jd_profile = jd_parser.parse_jd(jd_text)
        jd_embedding = scorer.get_jd_embedding(jd_profile)
        
        # 3. Features & Honeypots
        features_list = []
        hp_results = []
        for cand in candidates:
            feats = feature_extractor.extract_features(cand)
            is_hp, conf = honeypot_detector.is_honeypot(cand)
            features_list.append(feats)
            hp_results.append((is_hp, conf))
            
        # 4. Batch Score
        scores_list = scorer.batch_score_all(features_list, jd_profile)
        
        # Apply honeypot penalty
        final_scores = []
        for score, (is_hard_hp, _) in zip(scores_list, hp_results):
            if is_hard_hp:
                final_scores.append(round(score * 0.05, 4))
            else:
                final_scores.append(score)
                
        # Combine and Sort
        scored_candidates = []
        for cand, feat, score, (is_hard_hp, conf) in zip(candidates, features_list, final_scores, hp_results):
            scored_candidates.append({
                "candidate_id": feat["candidate_id"],
                "score": score,
                "features": feat,
                "candidate": cand,
                "is_hard_hp": is_hard_hp
            })
            
        # Tiebreak sort: score descending, then candidate_id ascending
        scored_candidates.sort(key=lambda x: str(x["candidate_id"]))
        scored_candidates.sort(key=lambda x: float(x["score"]), reverse=True)
        
        # Generate reasoning with deduplication helper
        seen_reasonings = set()
        for rank, item in enumerate(scored_candidates, 1):
            item["rank"] = rank
            reason = reasoning_generator.generate_reasoning(item["features"], item["score"], rank)
            
            # Deduplicate logic
            if reason in seen_reasonings:
                suffix = f" [{item['candidate_id']}]"
                if len(reason) + len(suffix) <= 200:
                    reason = reason.rstrip(".") + suffix + "."
                else:
                    reason = reason[:200 - len(suffix)] + suffix
            seen_reasonings.add(reason)
            item["reasoning"] = reason
            
        # Print top 10 with score and reasoning
        print("\n" + "="*80)
        print("TOP 10 CANDIDATES")
        print("="*80)
        for item in scored_candidates[:10]:
            print(f"Rank {item['rank']}: ID={item['candidate_id']} | Score={item['score']:.4f} | Reasoning: {item['reasoning']}")
        print("="*80 + "\n")
        
        # Checks
        # Check 1: All scores between 0.0 and 1.0
        scores_in_range = all(0.0 <= item["score"] <= 1.0 for item in scored_candidates)
        
        # Check 2: Uniqueness of reasoning
        all_reasonings = [item["reasoning"] for item in scored_candidates]
        reasonings_unique = len(all_reasonings) == len(set(all_reasonings))
        
        # Check 3: Scores decrease monotonically
        scores_monotonic = True
        for i in range(len(scored_candidates) - 1):
            if scored_candidates[i]["score"] < scored_candidates[i+1]["score"]:
                scores_monotonic = False
                break
                
        # Print results
        print("TEST RESULTS:")
        print(f"Check 1 (Scores between 0 and 1):     {'PASS' if scores_in_range else 'FAIL'}")
        print(f"Check 2 (Uniqueness of reasoning):    {'PASS' if reasonings_unique else 'FAIL'}")
        print(f"Check 3 (Monotonic score decrease):   {'PASS' if scores_monotonic else 'FAIL'}")
        
    except Exception as e:
        print(f"Execution failed with error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
